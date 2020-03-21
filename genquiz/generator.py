import docx
from docx.shared import Inches
from lxml import etree
from sympy.printing.mathml import mathml
import sympy
import numpy as np
import pandas as pd
import random
import pickle
import os


class Quiz:
    def __init__(self, student_id, student_name, questions, solutions, question_path, solution_path):
        # for correct Greek or special characters in MathML
        self.specialchar = {'&InvisibleTimes;': '&#8290;',
                            '&pi;': '&#960;',
                            }

        self.doc = docx.Document()

        self.student_id = student_id
        self.student_name = student_name
        self.solutions = solutions

        for i, question in enumerate(questions):
            self.__add_studentname()
            for item in question:
                self.__add(item)
            if i < len(questions) - 1:
                self.__add_page_break()
        self.__save(question_path, solution_path)

    def __add(self, item):
        if isinstance(item, Text):
            self.__add_text(item)
        if isinstance(item, Table):
            self.__add_table(item)
        if isinstance(item, Equation):
            self.__add_equation(item)
        if isinstance(item, Figure):
            self.__add_figure(item)

    def __math_to_word(self, eq, equal=None):
        math_ml = mathml(eq, printer='presentation')
        equal = mathml(equal, printer='presentation')
        # Creates mathml string
        if equal is None:
            mathml_string = '''
                <math xmlns="http://www.w3.org/1998/Math/MathML">
                    {}
                </math>
                '''.format(math_ml)
        else:
            mathml_string = '''
                <math xmlns="http://www.w3.org/1998/Math/MathML">
                    {0}
                    <mo>=</mo>
                    {1}
                 </math>
                '''.format(math_ml, equal)
        # Converts mathml string
        for ch in self.specialchar:
            mathml_string = mathml_string.replace(ch, self.specialchar[ch])

        tree = etree.fromstring(mathml_string)

        xslt = etree.parse(os.path.join(os.path.split(__file__)[0], 'MML2OMML.XSL'))
        transform = etree.XSLT(xslt)
        new_dom = transform(tree)
        return new_dom.getroot()

    def __add_table(self, obj):
        # header
        table = self.doc.add_table(rows=1, cols=len(obj.df.columns))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(obj.df.columns):
            hdr_cells[i].text = h
        # data
        for row in obj.df.values:
            row_cells = table.add_row().cells
            for i, e in enumerate(row):
                row_cells[i].text = str(e)
        table.style = 'Table Grid'

    def __add_figure(self, obj):
        if obj.img is not None:
            self.doc.add_picture(obj.img, width=obj.width)

    def __add_equation(self, obj):
        p = self.doc.add_paragraph()
        p._element.append(self.__math_to_word(obj.left, equal=obj.right))

    def __add_text(self, obj):
        self.doc.add_paragraph(obj.text)

    def __add_studentname(self):
        self.doc.add_heading(str(self.student_id) + '  ' + str(self.student_name), 0)

    def __add_page_break(self):
        self.doc.add_page_break()

    def __add_solution(self, solution):
        self.solution.append(solution)

    def __save(self, question_path, solution_path):
        self.doc.save(os.path.join(question_path, self.student_id + '.docx'))
        pickle.dump(self.solutions, open(os.path.join(solution_path, self.student_id + '.pk'), 'wb'))


class Text:
    def __init__(self, text, random_list=[]):
        self.text = text
        self.random_list = {}
        for kw in random_list:
            self.random_list[kw] = random.choice(random_list[kw])
            self.text = self.text.replace('{' + kw + '}', self.random_list[kw])


class Table:
    def __init__(self, field_desc):
        self.df = pd.DataFrame(columns=field_desc.keys())
        for f in field_desc:
            self.df[f] = field_desc[f]


class Equation:
    def __init__(self, left, right=None):
        self.left = left
        self.right = right


class Figure:
    def __init__(self, path, width=1.25):
        self.path = path
        self.img = None
        self.width = Inches(width)
        while True:
            fn = random.choice(os.listdir(path))
            if fn.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
                self.img = os.path.join(path, fn)
                with open(os.path.join(path, fn.split('.')[-2] + '.txt')) as f:
                    self.sol = np.array([f.readline()])
                break


class Solution:
    def __init__(self, solution, excel_cells, score=1):
        self.solution = solution
        self.excel_cells = excel_cells
        self.score = score


class Student:
    def __init__(self, path):
        df = pd.read_excel(path)
        self.id = df.iloc[:, 0].values.astype(np.str)
        self.name = df.iloc[:, 1].values

    def __iter__(self):
        self.i = 0
        return self

    def __next__(self):
        if self.i < len(self.id):
            self.i += 1
            return self.id[self.i - 1], self.name[self.i - 1]
        else:
            raise StopIteration
